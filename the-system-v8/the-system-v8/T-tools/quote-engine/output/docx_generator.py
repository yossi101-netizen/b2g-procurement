"""
KritiKaal Quote Engine — DOCX Generator
Produces the Two-Format Quote:
  Page 1: Bottom Line (all output formats)
  Page 2: Granular Breakdown (Full Unbundled only)

Driven entirely by QuoteResult — no business logic here.
"""
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from engine.models import OutputFormat, QuoteCurrency, QuoteResult


# ─────────────────────────────────────────────────────────────────
# BRAND PALETTE — change here to reskin all output
# ─────────────────────────────────────────────────────────────────
class _Brand:
    # Primary hex strings (for XML cell shading)
    DARK_HEX       = "1A2942"   # Deep navy — header, total row
    GOLD_HEX       = "C4972F"   # Warm gold — accents
    SECTION_HEX    = "EBF0F5"   # Light blue-gray — section rows
    SUBTOTAL_HEX   = "D0D9E3"   # Medium — subtotal rows
    PRICE_BG_HEX   = "F5F8FA"   # Near-white — price box background
    WHITE_HEX      = "FFFFFF"

    # RGBColor objects (for python-docx font/paragraph APIs)
    DARK    = RGBColor(0x1A, 0x29, 0x42)
    GOLD    = RGBColor(0xC4, 0x97, 0x2F)
    WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
    TEXT    = RGBColor(0x2C, 0x2C, 0x2C)
    MUTED   = RGBColor(0x6B, 0x7B, 0x8D)
    LIGHT   = RGBColor(0xF5, 0xF8, 0xFA)


# ─────────────────────────────────────────────────────────────────
# DOCUMENT CONSTANTS
# ─────────────────────────────────────────────────────────────────
_PAGE_WIDTH_IN    = 8.5
_LEFT_MARGIN_IN   = 1.0
_RIGHT_MARGIN_IN  = 1.0
_CONTENT_WIDTH_IN = _PAGE_WIDTH_IN - _LEFT_MARGIN_IN - _RIGHT_MARGIN_IN   # 6.5"
_CONTENT_TWIPS    = int(_CONTENT_WIDTH_IN * 1440)   # 9360 twips

_FONT = "Calibri"


# ─────────────────────────────────────────────────────────────────
# XML HELPERS
# ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour via XML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading if present
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120) -> None:
    """Set internal cell padding in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_width(cell, width_twips: int) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_width(table, width_twips: int) -> None:
    tblPr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _no_border(table) -> None:
    """Remove all borders from a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _single_border_cell(cell, sides: list, color: str = "C0C8D0", size: int = 6) -> None:
    """Add single-line border to specified sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(size))
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _paragraph_border_bottom(para, color: str = "C0C8D0", size: int = 6) -> None:
    """Add a bottom border line to a paragraph (used as a divider)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_para_spacing(para, before: int = 0, after: int = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def _keep_together(para) -> None:
    """Prevent paragraph from splitting across pages."""
    pPr = para._p.get_or_add_pPr()
    kT = OxmlElement("w:keepLines")
    pPr.append(kT)


# ─────────────────────────────────────────────────────────────────
# FORMATTING HELPERS
# ─────────────────────────────────────────────────────────────────

def _fmt(amount: float, currency: QuoteCurrency) -> str:
    sym = currency.symbol()
    return f"{sym}{amount:,.2f}"


def _fmt_usd(amount: float) -> str:
    return f"${amount:,.2f}"


def _fmt_pct(rate: float) -> str:
    return f"{rate:.1%}"


def _run(para, text: str, bold: bool = False, italic: bool = False,
         size_pt: float = 10, color: Optional[RGBColor] = None,
         font: str = _FONT) -> None:
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size_pt)
    if color:
        run.font.color.rgb = color


# ─────────────────────────────────────────────────────────────────
# MAIN GENERATOR CLASS
# ─────────────────────────────────────────────────────────────────

class QuoteDocxGenerator:
    """
    Generate a Two-Format Quote DOCX from a QuoteResult.
    Returns bytes — pass directly to Streamlit st.download_button().
    """

    def generate(self, result: QuoteResult) -> bytes:
        doc = Document()
        self._setup_document(doc)
        self._add_page1(doc, result)

        if result.inputs.output_format == OutputFormat.FULL_UNBUNDLED:
            doc.add_page_break()
            self._add_page2(doc, result)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ─────────────────────────────────────────────────────────────
    # DOCUMENT SETUP
    # ─────────────────────────────────────────────────────────────

    def _setup_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width    = Inches(_PAGE_WIDTH_IN)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(_LEFT_MARGIN_IN)
        section.right_margin  = Inches(_RIGHT_MARGIN_IN)
        section.top_margin    = Inches(0.6)
        section.bottom_margin = Inches(0.6)

        normal = doc.styles["Normal"]
        normal.font.name  = _FONT
        normal.font.size  = Pt(10)
        normal.font.color.rgb = _Brand.TEXT

    # ─────────────────────────────────────────────────────────────
    # PAGE 1 — BOTTOM LINE
    # ─────────────────────────────────────────────────────────────

    def _add_page1(self, doc: Document, r: QuoteResult) -> None:
        self._add_brand_header(doc, r)
        self._add_spacer(doc, 60)
        self._add_quote_meta(doc, r)
        self._add_divider(doc)
        self._add_destination_line(doc, r)
        self._add_spacer(doc, 80)
        self._add_price_box(doc, r)
        self._add_spacer(doc, 120)
        self._add_inclusions(doc, r)
        self._add_spacer(doc, 100)
        self._add_lead_time_block(doc, r)
        self._add_spacer(doc, 80)
        self._add_payment_terms(doc, r)
        self._add_page1_footer(doc, r)

    def _add_brand_header(self, doc: Document, r: QuoteResult) -> None:
        """Full-width dark header band with company name and contact."""
        table = doc.add_table(rows=1, cols=2)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)

        left_w  = int(_CONTENT_TWIPS * 0.62)
        right_w = _CONTENT_TWIPS - left_w

        left  = table.rows[0].cells[0]
        right = table.rows[0].cells[1]

        for cell, width in [(left, left_w), (right, right_w)]:
            _set_cell_bg(cell, _Brand.DARK_HEX)
            _set_cell_margins(cell, top=140, bottom=140, left=200, right=120)
            _set_cell_width(cell, width)

        # Left: company name + tagline
        p_name = left.paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p_name, before=0, after=40)
        _run(p_name, "KRITIKAAL", bold=True, size_pt=18, color=_Brand.WHITE)

        p_tag = left.add_paragraph()
        p_tag.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p_tag, before=0, after=0)
        _run(p_tag, "Premium India Leather Manufacturing", size_pt=8.5, color=_Brand.GOLD)

        # Right: contact details (right-aligned)
        p_contact = right.paragraphs[0]
        p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_spacing(p_contact, before=20, after=0)
        _run(p_contact, "contact@kritikaal.com\n", size_pt=8.5, color=_Brand.WHITE)
        _run(p_contact, "+972 58 620 8208\n", size_pt=8.5, color=_Brand.WHITE)
        _run(p_contact, "kritikaal.com", size_pt=8.5, color=_Brand.GOLD)

    def _add_quote_meta(self, doc: Document, r: QuoteResult) -> None:
        """2-column metadata table: client / ref / product / date / qty / validity."""
        table = doc.add_table(rows=3, cols=2)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)

        half = _CONTENT_TWIPS // 2
        data = [
            ("Client",   r.inputs.client_name,   "Quote Ref", r.quote_ref),
            ("Product",  r.inputs.product_ref,    "Date",      r.generated_at),
            ("Quantity", f"{r.inputs.units:,} units",
             "Valid for", "14 days from issue date"),
        ]
        for row_idx, (l_label, l_val, r_label, r_val) in enumerate(data):
            cells = table.rows[row_idx].cells
            for i, cell in enumerate(cells):
                _set_cell_width(cell, half)
                _set_cell_margins(cell, top=30, bottom=30, left=0, right=60)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _set_para_spacing(p, before=0, after=0)
                label = l_label if i == 0 else r_label
                val   = l_val   if i == 0 else r_val
                _run(p, f"{label}:  ", bold=True, size_pt=9, color=_Brand.MUTED)
                _run(p, val, size_pt=9, color=_Brand.TEXT)

    def _add_divider(self, doc: Document) -> None:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=80, after=80)
        _paragraph_border_bottom(p, color=_Brand.GOLD_HEX, size=12)

    def _add_spacer(self, doc: Document, after: int = 80) -> None:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=after)

    def _add_destination_line(self, doc: Document, r: QuoteResult) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p, before=0, after=40)
        _run(p, "DESTINATION:  ", bold=True, size_pt=9.5, color=_Brand.DARK)
        dest_text = f"{r.destination_port}  —  Delivered basis (DDP)"
        _run(p, dest_text, size_pt=9.5, color=_Brand.TEXT)

    def _add_price_box(self, doc: Document, r: QuoteResult) -> None:
        """Centered price highlight: per-unit | total order value."""
        currency = r.inputs.quote_currency

        table = doc.add_table(rows=2, cols=2)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)

        half = _CONTENT_TWIPS // 2

        # Row 0: labels
        for idx, label in enumerate(["PRICE PER UNIT", "TOTAL ORDER VALUE"]):
            cell = table.rows[0].cells[idx]
            _set_cell_bg(cell, _Brand.SECTION_HEX)
            _set_cell_width(cell, half)
            _set_cell_margins(cell, top=100, bottom=20, left=160, right=160)
            _single_border_cell(cell, ["top", "left" if idx == 0 else "right"],
                                 color=_Brand.DARK_HEX, size=12)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=0, after=0)
            _run(p, label, bold=True, size_pt=8, color=_Brand.MUTED)

        # Row 1: values
        values = [
            _fmt(r.per_unit_quote_currency, currency),
            _fmt(r.total_quote_currency, currency),
        ]
        for idx, val in enumerate(values):
            cell = table.rows[1].cells[idx]
            _set_cell_bg(cell, _Brand.PRICE_BG_HEX)
            _set_cell_width(cell, half)
            _set_cell_margins(cell, top=20, bottom=100, left=160, right=160)
            _single_border_cell(cell, ["bottom", "left" if idx == 0 else "right"],
                                 color=_Brand.DARK_HEX, size=12)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=0, after=0)
            _run(p, val, bold=True, size_pt=26, color=_Brand.DARK)

        # Inner vertical divider between the two price cells
        for row_idx in range(2):
            left_cell = table.rows[row_idx].cells[0]
            _single_border_cell(left_cell, ["right"], color=_Brand.DARK_HEX, size=8)

        # Note below box if currency not USD
        if currency != QuoteCurrency.USD:
            p_note = doc.add_paragraph()
            _set_para_spacing(p_note, before=40, after=0)
            p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p_note,
                 f"({currency.value} at {r.fx_rate_used:.4f} — includes {r.fx_buffer_rate:.0%} FX buffer)",
                 size_pt=8, color=_Brand.MUTED, italic=True)

    def _add_inclusions(self, doc: Document, r: QuoteResult) -> None:
        """What's included checklist."""
        p_head = doc.add_paragraph()
        _set_para_spacing(p_head, before=0, after=60)
        _run(p_head, "WHAT'S INCLUDED", bold=True, size_pt=9, color=_Brand.DARK)
        _paragraph_border_bottom(p_head, color="DEE3EA", size=4)

        qup_tier = r.inputs.qup_tier.value
        qup_labels = {
            "basic":    "AQL 2.5 quality inspection — 8-criteria evaluation, full results documentation",
            "standard": "AQL 2.5 quality inspection + production lead time guarantee (within 14 days of agreed ship date)",
            "maximum":  "AQL 2.5 inspection + lead time guarantee + material specification guarantee — full KritiKaal liability",
        }

        items = [
            f"Full-grain leather manufacturing to {r.inputs.product_ref} specification",
            qup_labels[qup_tier],
            "Sea freight from India to " + r.destination_port,
            "Customs clearance, import duty, and port handling",
            "Compliance documentation (Certificate of Origin, REACH, Azo Dye testing)",
            "KritiKaal production management, factory oversight, and AQL evaluation",
        ]
        if r.inputs.rex_certified and r.inputs.destination.value == "UK":
            items.insert(3, "UK DCTS Standard Preferences applied — 0% import duty (REX certified)")

        for item in items:
            p = doc.add_paragraph()
            _set_para_spacing(p, before=20, after=20)
            _run(p, "✓  ", bold=True, size_pt=9.5, color=_Brand.GOLD)
            _run(p, item, size_pt=9.5, color=_Brand.TEXT)

    def _add_lead_time_block(self, doc: Document, r: QuoteResult) -> None:
        table = doc.add_table(rows=1, cols=2)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)
        half = _CONTENT_TWIPS // 2

        pairs = [
            ("PRODUCTION LEAD TIME", r.production_lead_time_str),
            ("ESTIMATED DOOR-TO-DOOR", r.total_lead_time_str),
        ]
        for idx, (label, val) in enumerate(pairs):
            cell = table.rows[0].cells[idx]
            _set_cell_bg(cell, _Brand.SECTION_HEX)
            _set_cell_width(cell, half)
            _set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_para_spacing(p, before=0, after=20)
            _run(p, label + "\n", bold=True, size_pt=8, color=_Brand.MUTED)
            _run(p, val, bold=True, size_pt=10, color=_Brand.DARK)

    def _add_payment_terms(self, doc: Document, r: QuoteResult) -> None:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=40)
        _run(p, "PAYMENT TERMS:  ", bold=True, size_pt=9, color=_Brand.DARK)
        _run(p, "30% deposit at order confirmation — 70% balance prior to shipment.",
             size_pt=9, color=_Brand.TEXT)

    def _add_page1_footer(self, doc: Document, r: QuoteResult) -> None:
        p_div = doc.add_paragraph()
        _set_para_spacing(p_div, before=120, after=60)
        _paragraph_border_bottom(p_div, color="DEE3EA", size=4)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=0, after=0)
        _run(p, f"Quote Ref: {r.quote_ref}  |  Valid 14 days from {r.generated_at}  |  "
                f"Config: {r.config_version}",
             size_pt=7.5, color=_Brand.MUTED, italic=True)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p2, before=20, after=0)
        _run(p2, "KritiKaal — contact@kritikaal.com — +972 58 620 8208 — kritikaal.com",
             size_pt=7.5, color=_Brand.MUTED)

    # ─────────────────────────────────────────────────────────────
    # PAGE 2 — GRANULAR BREAKDOWN (Type A / Type C only)
    # ─────────────────────────────────────────────────────────────

    def _add_page2(self, doc: Document, r: QuoteResult) -> None:
        self._add_p2_header(doc, r)
        self._add_spacer(doc, 80)
        self._add_breakdown_table(doc, r)
        self._add_p2_footer(doc, r)

    def _add_p2_header(self, doc: Document, r: QuoteResult) -> None:
        """Compact header for Page 2."""
        table = doc.add_table(rows=1, cols=2)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)
        left_w  = int(_CONTENT_TWIPS * 0.60)
        right_w = _CONTENT_TWIPS - left_w

        left  = table.rows[0].cells[0]
        right = table.rows[0].cells[1]

        for cell, width in [(left, left_w), (right, right_w)]:
            _set_cell_bg(cell, _Brand.DARK_HEX)
            _set_cell_margins(cell, top=100, bottom=100, left=200, right=120)
            _set_cell_width(cell, width)

        p_left = left.paragraphs[0]
        _run(p_left, "COST BREAKDOWN  ", bold=True, size_pt=13, color=_Brand.WHITE)
        _run(p_left, f"— {r.quote_ref}", size_pt=10, color=_Brand.GOLD)

        p_right = right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p_right, "For reference and internal audit use.\nAll amounts in USD unless stated.",
             size_pt=8, color=_Brand.WHITE, italic=True)

    def _add_breakdown_table(self, doc: Document, r: QuoteResult) -> None:
        """
        3-column breakdown table:
        | Description                          | USD Total | Per Unit |
        """
        currency = r.inputs.quote_currency
        units    = r.inputs.units

        # Column widths: description 60%, total 22%, per-unit 18%
        col_w = [
            int(_CONTENT_TWIPS * 0.60),
            int(_CONTENT_TWIPS * 0.22),
            int(_CONTENT_TWIPS * 0.18),
        ]

        table = doc.add_table(rows=0, cols=3)
        _set_table_width(table, _CONTENT_TWIPS)
        _no_border(table)

        def add_header_row(label: str, col_widths: list) -> None:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                _set_cell_bg(cell, _Brand.DARK_HEX)
                _set_cell_width(cell, col_widths[i])
                _set_cell_margins(cell, top=80, bottom=80,
                                  left=160 if i == 0 else 80, right=80)
            p = row.cells[0].paragraphs[0]
            _run(p, label, bold=True, size_pt=9, color=_Brand.WHITE)
            for i in [1, 2]:
                hdr = ["TOTAL (USD)", "PER UNIT"][i - 1]
                ph = row.cells[i].paragraphs[0]
                ph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(ph, hdr, bold=True, size_pt=8, color=_Brand.WHITE)

        def add_data_row(desc: str, total_usd: float, per_unit_usd: float,
                          bold: bool = False, note: str = None,
                          bg: str = _Brand.WHITE_HEX,
                          indent: bool = True) -> None:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                _set_cell_bg(cell, bg)
                _set_cell_width(cell, col_w[i])
                _set_cell_margins(cell, top=55, bottom=55,
                                  left=160 if i == 0 else 80, right=80)
                _single_border_cell(cell, ["bottom"], color="E8EDF2", size=4)

            # Description cell
            p0 = row.cells[0].paragraphs[0]
            prefix = "    " if indent else ""
            _run(p0, prefix + desc, bold=bold, size_pt=9.5, color=_Brand.TEXT)
            if note:
                p_note = row.cells[0].add_paragraph()
                _set_para_spacing(p_note, before=0, after=0)
                _run(p_note, "    " + note, size_pt=7.5, color=_Brand.MUTED, italic=True)

            # Amount cells
            for i, val in enumerate([total_usd, per_unit_usd]):
                p = row.cells[i + 1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(p, _fmt_usd(val), bold=bold, size_pt=9.5, color=_Brand.TEXT)

        def add_subtotal_row(label: str, total_usd: float) -> None:
            row = table.add_row()
            per_unit = total_usd / units
            for i, cell in enumerate(row.cells):
                _set_cell_bg(cell, _Brand.SUBTOTAL_HEX)
                _set_cell_width(cell, col_w[i])
                _set_cell_margins(cell, top=70, bottom=70,
                                  left=160 if i == 0 else 80, right=80)
            p0 = row.cells[0].paragraphs[0]
            _run(p0, label, bold=True, size_pt=9, color=_Brand.DARK)
            for i, val in enumerate([total_usd, per_unit]):
                p = row.cells[i + 1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(p, _fmt_usd(val), bold=True, size_pt=9, color=_Brand.DARK)

        def add_section_spacer() -> None:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                _set_cell_bg(cell, _Brand.WHITE_HEX)
                _set_cell_width(cell, col_w[i])
                _set_cell_margins(cell, top=20, bottom=20, left=0, right=0)

        # ── Column header row ──────────────────────────────────
        add_header_row("SECTION 1 — MANUFACTURING PASSTHROUGH", col_w)

        add_data_row(
            f"Factory FOB — {r.inputs.product_ref} ({units:,} units @ ${r.factory_fob_per_unit:,.2f}/unit)",
            r.factory_fob_total,
            r.factory_fob_per_unit,
            indent=False,
        )
        add_data_row(
            f"Product packaging ({r.inputs.packaging_type.display()})",
            r.packaging_total,
            r.packaging_per_unit,
        )
        add_subtotal_row("Subtotal — Manufacturing Passthrough", r.manufacturing_passthrough)

        add_section_spacer()

        # ── Section 2: Logistics ───────────────────────────────
        add_header_row("SECTION 2 — LOGISTICS & CUSTOMS", col_w)

        add_data_row(
            f"Sea freight ({r.freight_mode_display}, India → {r.destination_port})",
            r.freight_total,
            r.freight_total / units,
            indent=False,
        )
        duty_note = None
        if r.duty_total == 0:
            duty_note = "REX certification applied — 0% duty confirmed"
        add_data_row(
            f"Import duty — {r.duty_display}",
            r.duty_total,
            r.duty_total / units,
            note=duty_note,
        )
        add_data_row("Customs clearance (broker fee)",   r.broker_total,   r.broker_total / units)
        add_data_row("Port handling",                    r.port_total,     r.port_total / units)
        add_data_row("Marine insurance (0.5% of CIF)",  r.insurance_total, r.insurance_total / units)
        add_subtotal_row("Subtotal — Logistics & Customs", r.logistics_passthrough)

        add_section_spacer()

        # ── Section 3: KritiKaal Services ─────────────────────
        add_header_row("SECTION 3 — KRITIKAAL SERVICES", col_w)

        add_data_row(
            "Production Management & QA Services",
            r.production_management_qa,
            r.production_management_qa / units,
            bold=False,
            indent=False,
            note=(
                "Covers: sample coordination, pre-production management, "
                "compliance documentation, AQL 2.5 quality oversight, "
                "and managed manufacturing service fee."
            ),
        )
        add_data_row(
            f"Quality Underwriting Premium — {r.inputs.qup_tier.display()} ({_fmt_pct(r.qup_rate)})",
            r.qup_total,
            r.qup_total / units,
            note=r.qup_description,
        )

        qup_subtotal = r.production_management_qa + r.qup_total
        add_subtotal_row("Subtotal — KritiKaal Services", qup_subtotal)

        add_section_spacer()

        # ── Section 4: Currency Management ───────────────────
        if r.fx_buffer_amount > 0:
            add_header_row("SECTION 4 — CURRENCY MANAGEMENT", col_w)
            add_data_row(
                f"FX Buffer ({_fmt_pct(r.fx_buffer_rate)}, USD → {r.inputs.quote_currency.value})",
                r.fx_buffer_amount,
                r.fx_buffer_amount / units,
                indent=False,
                note="Hedges currency movement between quotation and settlement.",
            )
            add_section_spacer()

        # ── Grand Total ────────────────────────────────────────
        # USD total row
        row_usd = table.add_row()
        for i, cell in enumerate(row_usd.cells):
            _set_cell_bg(cell, _Brand.DARK_HEX)
            _set_cell_width(cell, col_w[i])
            _set_cell_margins(cell, top=100, bottom=100,
                              left=160 if i == 0 else 80, right=80)

        p_usd = row_usd.cells[0].paragraphs[0]
        _run(p_usd, "TOTAL ORDER VALUE (USD)", bold=True, size_pt=10, color=_Brand.WHITE)

        for i, val in enumerate([r.total_usd, r.total_usd / units]):
            p = row_usd.cells[i + 1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(p, _fmt_usd(val), bold=True, size_pt=10, color=_Brand.WHITE)

        # Converted currency total row (if not USD)
        if r.inputs.quote_currency != QuoteCurrency.USD:
            row_fx = table.add_row()
            currency = r.inputs.quote_currency
            for i, cell in enumerate(row_fx.cells):
                _set_cell_bg(cell, _Brand.GOLD_HEX)
                _set_cell_width(cell, col_w[i])
                _set_cell_margins(cell, top=100, bottom=100,
                                  left=160 if i == 0 else 80, right=80)

            p_fx = row_fx.cells[0].paragraphs[0]
            _run(p_fx,
                 f"TOTAL ORDER VALUE ({currency.value} at {r.fx_rate_used:.4f})",
                 bold=True, size_pt=10, color=_Brand.WHITE)

            for i, val in enumerate([r.total_quote_currency, r.per_unit_quote_currency]):
                p = row_fx.cells[i + 1].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(p, _fmt(val, currency), bold=True, size_pt=10, color=_Brand.WHITE)

    def _add_p2_footer(self, doc: Document, r: QuoteResult) -> None:
        self._add_spacer(doc, 80)
        p = doc.add_paragraph()
        _paragraph_border_bottom(p, color="DEE3EA", size=4)
        _set_para_spacing(p, before=0, after=60)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p2, before=0, after=20)
        _run(p2,
             "All costs verified at time of quotation. Freight rates valid 14 days from issue date. "
             "Currency conversion applies published buffer rates. Duty rates subject to "
             "change — verify current tariff schedule before order placement.",
             size_pt=7.5, color=_Brand.MUTED, italic=True)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p3, before=0, after=0)
        _run(p3,
             f"VAT/Maam is not included in this quote and is recoverable by VAT-registered businesses. "
             f"Config version: {r.config_version}.",
             size_pt=7.5, color=_Brand.MUTED, italic=True)
